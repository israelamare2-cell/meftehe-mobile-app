import streamlit as st
import os
import google.generativeai as genai
from docx import Document
from docx.enum.section import WD_ORIENT
import io
import requests
import sqlite3
import hashlib
import time
import re

def format_math_for_word(text):
    """LaTeX ኮዶችን ወደ ፅድት ያለና ወርድ ላይ ወደሚነበብ የሒሳብ ምልክት መቀየሪያ"""
    if not text:
        return ""

    # 1. መጀመሪያ የላቴክስ ምልክቶችን ($, $$, \[, \]) እናጠፋለን
    text = text.replace('$$', '').replace('$', '').replace(r'\[', '').replace(r'\]', '')

    # 2. Fractions (ክፍልፋይ): \frac{a}{b} -> (a/b)
    text = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', r'(\1/\2)', text)

    # 3. Powers (ስኩዌር): x^{2} -> x²
    superscripts = {"0":"⁰","1":"¹","2":"²","3":"³","4":"⁴","5":"⁵","6":"⁶","7":"⁷","8":"⁸","9":"⁹","n":"ⁿ","x":"ˣ"}
    def replace_power(match):
        p = match.group(1)
        return "".join(superscripts.get(c, "^"+c) for c in p)
    text = re.sub(r'\^\{?([0-9nx]+)\}?', replace_power, text)

    # 4. Roots (ስኩዌር ሩት): \sqrt{x} -> √x
    text = re.sub(r'\\sqrt\{(.+?)\}', r'√(\1)', text)

    # 5. የሒሳብ ምልክቶች (Symbols)
    symbols = {
        r'\times': '×', r'\div': '÷', r'\pm': '±', r'\neq': '≠',
        r'\pi': 'π', r'\degree': '°', r'\therefore': '∴', r'\le': '≤', r'\ge': '≥'
    }
    for latex, symbol in symbols.items():
        text = text.replace(latex, symbol)

    # 6. አላስፈላጊ የሆኑ የላቴክስ ትዕዛዞችን ማጽዳት
    text = text.replace(r'\left', '').replace(r'\right', '').replace(r'\{', '{').replace(r'\}', '}')
    return text
# --- 1. የገጽ ቅንብር ---
st.set_page_config(page_title="Meftehe AI App", page_icon="📖", layout="wide")

# --- 2. ኮንፊገሬሽን እና API Keys ---
GITHUB_USER = "israelamare2-cell"
GITHUB_REPO = "meftehe-bot"
RELEASE_TAG = "v1" 
GITHUB_BASE_URL = f"https://github.com/{GITHUB_USER}/{GITHUB_REPO}/releases/download/{RELEASE_TAG}/"

# API Load Balancing (ልክ እንደ ቦቱ)
GEMINI_API_KEYS_STR = os.getenv('GEMINI_API_KEYS', os.getenv('GEMINI_API_KEY', ''))
API_KEY_LIST = [k.strip() for k in GEMINI_API_KEYS_STR.split(',')] if GEMINI_API_KEYS_STR else []

# --- 3. ዳታቤዝ እና SMART CACHE (ከቦቱ የተወሰደ) ---
def init_db():
    conn = sqlite3.connect('meftehe_national_data.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS gemini_cache 
                 (prompt_hash TEXT PRIMARY KEY, response_text TEXT)''')
    conn.commit()
    conn.close()

def get_cached_response(prompt_text, file_bytes):
    file_hash = hashlib.md5(file_bytes).hexdigest()
    combined = prompt_text + file_hash
    hash_val = hashlib.md5(combined.encode()).hexdigest()
    conn = sqlite3.connect('meftehe_national_data.db', check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT response_text FROM gemini_cache WHERE prompt_hash=?", (hash_val,))
    result = c.fetchone()
    conn.close()
    return result[0] if result else None, hash_val

def save_to_cache(hash_val, response_text):
    conn = sqlite3.connect('meftehe_national_data.db', check_same_thread=False)
    c = conn.cursor()
    try:
        c.execute("INSERT OR REPLACE INTO gemini_cache VALUES (?, ?)", (hash_val, response_text))
        conn.commit()
    except:
        pass
    conn.close()

init_db()

# --- 4. መጽሐፍ ከ GitHub የሚያወርድ (ከቦቱ የተወሰደ) ---
@st.cache_data
def get_book(grade, subject):
    folder = "books"
    if not os.path.exists(folder): os.makedirs(folder)
    
    file_name = f"grade{grade}_{subject.lower().replace(' ', '_')}.pdf"
    path = os.path.join(folder, file_name)
    
    if os.path.exists(path): return path
    
    url = f"{GITHUB_BASE_URL}{file_name}"
    try:
        r = requests.get(url, stream=True)
        if r.status_code == 200:
            with open(path, 'wb') as f:
                for chunk in r.iter_content(8192): f.write(chunk)
            return path
    except: return None
    return None

# --- 5. መሠረታዊ ዳታዎች ---
ALL_SUBJECTS = ["Mathematics", "Physics", "Chemistry", "Biology", "General Science", "English", "Social Studies", "Citizenship", "Amharic", "Afaan Oromoo", "Environmental Science", "Moral Education", "PVA", "HPE", "CTE", "Agriculture", "Economics", "IT"]
ALL_ASSESSMENT_TYPES = ["Mid Exam", "Final Exam", "Worksheet", "Quiz", "National Prep", "Model Exam", "Test"]

# --- 6. የአፑ ፊት (User Interface) ---
st.sidebar.title("🌟 መፍትሔ (Meftehe)")
st.sidebar.caption("የመምህራን ረዳት ሞባይል አፕ")

# ዋና ምርጫዎች
app_lang = st.sidebar.selectbox("🌍 ቋንቋ / Language", ["am", "or", "ti", "so", "en"])
mode = st.sidebar.radio("🛠 ምን ማዘጋጀት ይፈልጋሉ?", ["exam", "note", "lesson", "review"], format_func=lambda x: {"exam": "📝 የፈተና ዝግጅት", "note": "📚 የማስተማሪያ ኖት", "lesson": "📅 ዕለታዊ የትምህርት ዕቅድ", "review": "🔍 የመፅሀፍ ግምገማ"}[x])

st.title({"exam": "📝 የፈተና ዝግጅት", "note": "📚 የማስተማሪያ ኖት", "lesson": "📅 ዕለታዊ የትምህርት ዕቅድ (SMASE)", "review": "🔍 የመፅሀፍ ግምገማ እና ኦዲት"}[mode])

col1, col2 = st.columns(2)
with col1:
    subj = st.selectbox("📚 ትምህርት ይምረጡ", ALL_SUBJECTS)
with col2:
    grd = st.number_input("📖 የክፍል ደረጃ", 1, 12, 7)

chp = st.text_input("📂 ምዕራፍ ይምረጡ (ለምሳሌ፦ Chapter 2 ወይም All)")

# እንደየ ምርጫው (Mode) የሚቀያየሩ ቅፆች
tos_config = "auto"
num_sets = 1
lang_output_option = None

if mode == "exam":
    col3, col4, col5 = st.columns(3)
    with col3: as_type = st.selectbox("📝 የፈተና አይነት", ALL_ASSESSMENT_TYPES)
    with col4: diff = st.selectbox("📊 የክብደት ደረጃ", ["Easy", "Medium", "Hard", "MixedFair"])
    with col5: bloom = st.selectbox("🧠 Bloom's Taxonomy", ["Knowledge", "Understanding", "Application", "Analysis", "Evaluation", "Creation", "Mixed"])
    
    is_language = subj.lower() in ["amharic", "english", "afaan oromoo"]
    if is_language:
        lang_output_option = st.radio("📖 ለቋንቋ ትምህርት ምን ይዘጋጅ?", ["PassageOnly", "QuestionOnly", "Both"], format_func=lambda x: {"PassageOnly": "📄 ምንባብ ብቻ", "QuestionOnly": "❓ ጥያቄ ብቻ", "Both": "📑 ምንባብ እና ጥያቄ"}[x])
    else:
        num_sets = st.selectbox("🛡️ የሴት ብዛት", [1, 2, 4])
    
    tos_config = st.text_input("🔢 የጥያቄ ብዛትና መዋቅር (ለምሳሌ፦ ምርጫ=10, እውነት/ሐሰት=5)", value="auto")

elif mode == "lesson":
    pg_type = st.radio("🔢 ገፅ እንዴት መምረጥ ይፈልጋሉ?", ["አንድ ገፅ ብቻ", "የገፅ ክልል (ከ-እስ)"])
    tos_config = st.text_input("🖋 የገፅ ቁጥሩን ይጻፉ", placeholder="ለምሳሌ: 12 ወይም 12-15")

elif mode == "review":
    review_type = st.selectbox("🧐 የግምገማ ዘርፍ ይምረጡ", ["FullAudit", "Pedagogy", "Assessment", "Indigenous", "21stCentury", "Inclusivity"])
    page_range = st.selectbox("📄 ለመገምገም የሚፈልጉትን የገፅ ክልል", ["1-20", "21-50", "51-100", "101-150", "151-200", "All Pages"])
    tos_config = st.text_input("🖋 ልዩ ትኩረት እንዲሰጥበት የሚፈልጉት ነጥብ (ወይም 'auto')")

elif mode == "note":
    note_style = st.selectbox("✨ የኖት አይነት ይምረጡ", ["Objectives", "Comprehensive", "Examples", "Summary", "ReviewQs", "FullPackage"])
    tos_config = st.text_input("📄 ማስታወሻው እንዲያካትት የሚፈልጉት ልዩ ነጥብ", value="auto")

# --- 7. የማመንጨት ሂደት (Generation) ---
if st.button("🚀 አዘጋጅ / Generate", use_container_width=True):
    if not chp or not tos_config:
        st.warning("እባክዎ የምዕራፍ ስም እና ሌሎች ባዶ ቦታዎችን ይሙሉ!")
    else:
        with st.spinner("🔍 መፅሀፉን እያወረድኩ እና AI እያዘጋጀ ነው (ይህ ጥቂት ደቂቃዎች ሊወስድ ይችላል)..."):
            book_path = get_book(grd, subj)
            
            if not book_path:
                st.error("❌ መፅሀፉ በ GitHub Release ላይ አልተገኘም!")
            else:
                # የቋንቋ ህጎች ከቦቱ
                lang_map = {
                    'am': "STRICTLY in AMHARIC language.",
                    'or': "STRICTLY in AFAAN OROMOO language using professional terms.",
                    'ti': "STRICTLY in TIGRIGNA language.",
                    'so': "STRICTLY in SOMALI language.",
                    'en': "STRICTLY in ENGLISH language."
                }
                
                target_subject = subj.lower()
                if target_subject == "afaan oromoo": lang_rule = lang_map['or']
                elif target_subject == "amharic": lang_rule = lang_map['am']
                elif target_subject == "english": lang_rule = lang_map['en']
                else: lang_rule = lang_map[app_lang]

                # Prompt አዘገጃጀት (ከቦቱ የተወሰደ)
                if mode == "lesson":
                    prompt = f"""You are a Professional Curriculum Developer specializing in SMASE (Active Learning).
                    TASK: Create a DAILY LESSON PLAN based on Chapter: {chp} and Page: {tos_config}.
                    STRICT REQUIREMENTS:
                    1. LANGUAGE: {lang_rule}
                    2. PEDAGOGY: Follow SMASE (Active Learning). Ensure it's learner-centered.
                    3. STYLE: Use VERY SHORT BULLET POINTS. NO long sentences.
                    4. FORMAT: Use a CLEAR TABLE for the Teacher/Student activity sections.
                    HEADER INFO:
                    - School: የካ ተራራ ቅድመ አንደኛ፣ አንደኛ እና መካከለኛ ደረጃ ትምህርት ቤት
                    - Teacher: እስራኤል አማረ
                    
                    SECTIONS TO INCLUDE (Short & Precise):
                    - Objectives (አላማዎች)
                    - Significance (አስፈላጊነት)
                    - Prior Knowledge (ቀዳሚ ዕዉቀት)
                    - Competency (አጥጋቢ የመማር ብቃት)
                    
                    TABLE STRUCTURE:
                    Generate a table with these columns: [የመማር ማስተማር ቅደም ተከተል, ክፍለ ጊዜ, ይዘት, የመምህሩ ተግባር, የተማሪ ተግባር, ምዘና, የመርጃ መሣሪያ]
                    
                    DIFFERENTIATED SUPPORT (Short examples):
                    - For High Achievers, Average Students, Low Achievers, Special Needs
                    """
                elif mode == "exam":
                    lang_output_instruction = ""
                    if lang_output_option == "PassageOnly": lang_output_instruction = "IMPORTANT: Provide ONLY the Reading Passage. Do NOT generate any questions."
                    elif lang_output_option == "QuestionOnly": lang_output_instruction = "IMPORTANT: Generate ONLY the questions. Do NOT include the reading passage."
                    elif lang_output_option == "Both": lang_output_instruction = "IMPORTANT: Provide the Reading Passage FIRST, followed by the related questions."

                    prompt = f"""You are an expert Ethiopian National Examiner.
                    STRICT COMPLIANCE:
                    1. SOURCE: Use ONLY the provided PDF. Focus on Chapter: {chp}, Bloom Level: {bloom}, Difficulty: {diff}.
                    2. USER COMMAND: Create exam structure: {tos_config}.
                    3. LANGUAGE: {lang_rule}
                    4. SYMBOLS: ALL formulas in LaTeX using $inline$ or $$display$$.
                    5. LANGUAGE SUBJECT RULE: {lang_output_instruction}
                    6. OUTPUT: {num_sets} different sets. Include TOS, Exam, and Answer Key."""
                elif mode == "review":
                    prompt = f"""You are a Precise Curriculum Auditor.
                    TASK: Conduct a PAGE-BY-PAGE Audit of the PDF for Page Range: {page_range} / Chapter: {chp}.
                    REVIEW SCOPE: {review_type}
                    
                    STRICT OUTPUT STRUCTURE:
                    1. EXECUTIVE SUMMARY
                    2. DETAILED PAGE-BY-PAGE FINDINGS
                    3. CRITICAL ERRORS TABLE
                    4. PEDAGOGICAL ALIGNMENT
                    FORMATTING: LANGUAGE: {lang_rule}. USER SPECIAL NOTE: {tos_config}"""
                else:
                    prompt = f"Professional Curriculum Expert Note Generation for Chapter {chp}... Style: {note_style}. Language: {lang_rule}..."

                # ፋይሉን ማንበብ እና AI መጥራት
                with open(book_path, "rb") as f: file_data = f.read()
                
                cached_response, prompt_hash = get_cached_response(prompt, file_data)
                
                if cached_response:
                    st.success("🚀 መረጃው ከዚህ ቀደም ስለተጠየቀ ከዳታቤዝ (Cache) በፍጥነት ተገኝቷል!")
                    raw_content = cached_response
                else:
                    max_retries = len(API_KEY_LIST) if API_KEY_LIST else 1
                    success = False
                    current_key_index = 0
                    
                    for attempt in range(max_retries * 2):
                        try:
                            if API_KEY_LIST: genai.configure(api_key=API_KEY_LIST[current_key_index % len(API_KEY_LIST)])
                            model = genai.GenerativeModel('gemini-2.5-flash')
                            response = model.generate_content([{"mime_type": "application/pdf", "data": file_data}, prompt])
                            raw_content = response.text.replace("###", "").replace("##", "")
                            save_to_cache(prompt_hash, raw_content)
                            success = True
                            st.success("✅ አዲስ መረጃ በተሳካ ሁኔታ ተዘጋጅቷል!")
                            break
                        except Exception as e:
                            current_key_index += 1
                            time.sleep(2)
                            continue
                    
                    if not success: st.error("⚠️ ይቅርታ፣ ሁሉም የ API ቁልፎች አሁን ላይ ተጨናንቀዋል። እባክዎ ከጥቂት ደቂቃዎች በኋላ ይሞክሩ።")

                # --- 8. ወደ Word መቀየር (ልክ እንደ ቦቱ) ---
                if 'raw_content' in locals():
                    with st.expander("👀 የተዘጋጀውን መረጃ እዚህ ይመልከቱ"):
                        st.write(raw_content)

                    doc = Document()
                    if mode == "lesson":
                        # Landscape ማረጋገጫ ለ Lesson Plan
                        section = doc.sections[0]
                        section.orientation = WD_ORIENT.LANDSCAPE
                        new_width, new_height = section.page_height, section.page_width
                        section.page_width, section.page_height = new_width, new_height
                        
                        header = doc.add_paragraph("ዕለታዊ የትምህርት ዕቅድ")
                        header.alignment = 1
                        info = doc.add_paragraph()
                        info.add_run(f"የመምህሩ ስም: እስራኤል አማረ\t\t\t\tየት/ቤቱ ስም: የካ ተራራ ቅድመ አንደኛ፣ አንደኛ እና መካከለኛ ደረጃ ትምህርት ቤት\n")
                        info.add_run(f"የትም ዓይነት: {subj}\t\t\t\tምዕራፍ: {chp}\n")
                        info.add_run(f"የክፍል ደረጃ: {grd}\t\t\t\tየዕለቱ ገፅ: {tos_config}")
                        doc.add_paragraph("\n" + raw_content)
                        doc.add_paragraph("\nመምህር: እስራኤል አማረ _________ \t የት/ክፍል ተጠሪ: አስመራወርቅ ሀይሌ _________ \t ም/ር/መ/ር: ከበደ ተስፋዪ _________")
                    else:
                        title = doc.add_heading(f"{subj} - Grade {grd} {mode.upper()}", 0)
                        title.alignment = 1 
                        sections = raw_content.split('\n\n')
                        for section in sections:
                            clean_sec = section.strip()
                            if not clean_sec: continue
                            if clean_sec.startswith("[Page") or clean_sec.startswith("Page") or ":" in clean_sec.split('\n')[0]:
                                p = doc.add_paragraph()
                                run = p.add_run(clean_sec)
                                run.bold = True
                            elif "|" in clean_sec: doc.add_paragraph(clean_sec) 
                            else: doc.add_paragraph(clean_sec)

                    target = io.BytesIO()
                    doc.save(target)
                    st.download_button("📥 በ Word ፋይል አውርድ (Download Document)", target.getvalue(), f"{subj}_{mode}.docx", type="primary", use_container_width=True)

st.divider()
st.caption("© 2026 Meftehe AI - Built for Ethiopian Teachers")
